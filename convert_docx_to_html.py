import docx
import sys
from html import escape

# 将 .docx 文件转换为 HTML
def docx_to_html(docx_path):
    doc = docx.Document(docx_path)
    html = ""

    # 遍历文档中的每个段落，转换为 <p> 标签
    for para in doc.paragraphs:
        html += f"<p>{escape(para.text)}</p>"

    # 获取表格并转换为 HTML 表格
    for table in doc.tables:
        html += "<table border='1'>"
        for row in table.rows:
            html += "<tr>"
            for cell in row.cells:
                html += f"<td>{escape(cell.text)}</td>"
            html += "</tr>"
        html += "</table>"

    return html

if __name__ == "__main__":
    # 获取命令行传入的文件路径
    docx_path = sys.argv[1]
    html_content = docx_to_html(docx_path)
    print(html_content)  # 输出 HTML 内容
